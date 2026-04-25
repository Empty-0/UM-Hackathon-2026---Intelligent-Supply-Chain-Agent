"""
Procurement Agents — Flask Web Application

Routes UI requests through the three-agent pipeline:
  Agent 1 (extract_requirements) → Agent 2 (source_all_items) → Agent 3 (generate_all_drafts)

Supports:
  - Raw text input (natural language)
  - Document upload: .docx (python-docx), .txt, .pdf (basic text), images (pytesseract OCR + GLM-4V visual analysis)
  - Multi-item processing with geographic risk alerts
  - UPDATE/CORRECTION intent handling
"""

import csv
import json
import os
import tempfile
import base64

from flask import Flask, render_template, request, jsonify
from zhipuai import ZhipuAI

from agents import (
    extract_requirements,
    source_all_items,
    generate_all_drafts,
    send_all_emails,
)

app = Flask(__name__)

CSV_PATH = os.path.join(os.path.dirname(os.path.abspath(__file__)), "suppliers.csv")

# ---------- 视觉模型分析函数 ----------
def analyze_image_with_glmv(file_path, api_key):
    """
    使用智谱 GLM-4V 模型分析图片内容，返回结构化 JSON。
    """
    try:
        client = ZhipuAI(api_key="a82eff1d7e5e04d8182563072a428d063.E4T49tzSmw5ezjii")
        with open(file_path, "rb") as image_file:
            image_base64 = base64.b64encode(image_file.read()).decode('utf-8')
        
        response = client.chat.completions.create(
            model="glm-4v-plus",  # 可用 glm-4v-plus 或 glm-4v
            messages=[{
                "role": "user",
                "content": [
                    {
                        "type": "image_url",
                        "image_url": {"url": f"data:image/jpeg;base64,{image_base64}"}
                    },
                    {
                        "type": "text",
                        "text": (
                            "请分析这张图片，你的回答必须是一个合法的 JSON 对象。\n"
                            "1. 描述图片中主要的物体是什么。\n"
                            "2. 识别图片中是否有与采购相关的文字，例如产品名、数量、价格等，并提取出来。\n"
                            "3. 根据图片内容，猜测用户的潜在采购需求。\n"
                            "请严格按照以下格式输出 JSON：\n"
                            "{\"description\": \"图片描述\", \"detected_objects\": [\"检测到的物品列表\"], "
                            "\"extracted_text\": \"提取的文字\", \"potential_requirement\": \"用户可能的需求\"}\n"
                            "如果没有提取到文字，extracted_text的值应为空字符串。"
                        )
                    }
                ]
            }],
            temperature=0.3
        )
        result_text = response.choices[0].message.content
        # 清理可能的 markdown 标记
        result_text = result_text.strip().strip('```json').strip('```').strip()
        result_json = json.loads(result_text)
        return result_json
    except Exception as e:
        print(f"视觉分析失败: {str(e)}")
        return {
            "description": "分析失败",
            "detected_objects": [],
            "extracted_text": "",
            "potential_requirement": ""
        }


# ── Document Extraction Helpers ────────────────────────────────────────

def _extract_docx(file_storage) -> str:
    """Extract text from a .docx file using python-docx."""
    try:
        from docx import Document
    except ImportError:
        return ""
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        file_storage.save(tmp.name)
        tmp_path = tmp.name
    try:
        doc = Document(tmp_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        tables_text = []
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                tables_text.append(" | ".join(cells))
        return "\n".join(paragraphs + tables_text).strip()
    finally:
        os.unlink(tmp_path)


def _extract_image(file_storage) -> str:
    """
    增强版图片处理：优先使用 GLM-4V 视觉模型理解图片内容并提取文字，
    如果模型未提取到文字则回退到 pytesseract OCR。
    """
    with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as tmp:
        file_storage.save(tmp.name)
        tmp_path = tmp.name
    try:
        # 获取 API Key（建议从环境变量读取，此处直接使用与 agents 相同的 key）
        # 注意：请将下面的字符串替换为你的实际 API Key
        API_KEY = "你的智谱API Key"   # <--- 替换为真实 Key
        
        # 1. 视觉模型分析
        vision_analysis = analyze_image_with_glmv(tmp_path, API_KEY)
        print(f"视觉模型分析结果：{vision_analysis}")
        
        # 2. 如果视觉模型提取到了文字，优先使用
        if vision_analysis.get("extracted_text"):
            text_content = vision_analysis.get("extracted_text")
        else:
            # 3. 后备 OCR
            print("视觉模型未提取到文字，启用 OCR 识别...")
            try:
                from PIL import Image
                import pytesseract
                img = Image.open(tmp_path)
                text_content = pytesseract.image_to_string(img).strip()
            except ImportError:
                text_content = ""
        
        # 4. 如果仍然没有文字，则基于模型对图像内容的描述生成采购需求提示
        if not text_content:
            description = vision_analysis.get("description", "")
            detected_objects = vision_analysis.get("detected_objects", [])
            potential_requirement = vision_analysis.get("potential_requirement", "")
            
            if description and detected_objects:
                product_names = ", ".join(detected_objects)
                text_content = f"采购需求：我需要采购 {product_names}。参考信息：{description}。{potential_requirement}"
            else:
                text_content = "上传了一张图片，需要基于图片内容进行采购。"
                
        return text_content.strip()
        
    except Exception as e:
        print(f"增强型图片处理失败: {e}")
        return ""
    finally:
        os.unlink(tmp_path)


def _extract_pdf_text(file_storage) -> str:
    """
    Extract text from a PDF file.
    Tries PyMuPDF (fitz) first, then pdfminer, then raw text fallback.
    """
    with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp:
        file_storage.save(tmp.name)
        tmp_path = tmp.name

    text = ""
    try:
        # Try PyMuPDF
        try:
            import fitz
            doc = fitz.open(tmp_path)
            for page in doc:
                text += page.get_text()
            doc.close()
            if text.strip():
                return text.strip()
        except ImportError:
            pass

        # Try pdfminer
        try:
            from pdfminer.high_level import extract_text as pdfminer_extract
            text = pdfminer_extract(tmp_path)
            if text.strip():
                return text.strip()
        except ImportError:
            pass

        # Raw text fallback
        with open(tmp_path, "rb") as f:
            raw = f.read()
        text = raw.decode("utf-8", errors="ignore")
        cleaned = text.replace("\x00", " ")
        cleaned = " ".join(cleaned.split())
        return cleaned.strip() if len(cleaned) > 10 else ""
    finally:
        os.unlink(tmp_path)


def _extract_text_from_upload(file_storage) -> tuple[str, str]:
    """
    Extract text from an uploaded file.
    Returns (extracted_text, source_type).
    """
    filename = file_storage.filename or ""

    if filename.endswith(".txt"):
        raw = file_storage.read()
        try:
            text = raw.decode("utf-8")
        except UnicodeDecodeError:
            text = raw.decode("latin-1")
        return text.strip(), "txt"

    if filename.endswith(".docx"):
        text = _extract_docx(file_storage)
        return text, "docx"

    if filename.endswith(".doc"):
        # .doc is old binary format — try python-docx (may fail)
        text = _extract_docx(file_storage)
        if not text:
            return "", "doc_unsupported"
        return text, "doc"

    if filename.lower().endswith((".png", ".jpg", ".jpeg", ".tiff", ".bmp", ".gif")):
        text = _extract_image(file_storage)
        return text, "image_vision_ocr"

    if filename.endswith(".pdf"):
        text = _extract_pdf_text(file_storage)
        return text, "pdf"

    return "", "unsupported"


# ── Pages ─────────────────────────────────────────────────────────────

@app.route("/")
def index():
    return render_template("index.html")


# ── Pipeline API ──────────────────────────────────────────────────────

def _run_pipeline_from_text(user_input: str, source: str = "text"):
    """Shared pipeline logic for both text and document inputs."""
    if not user_input:
        return jsonify({"error": "empty request"}), 400

    # Agent 1: Requirement Extraction
    req_json = extract_requirements(user_input)
    req_parsed = json.loads(req_json)

    if "error" in req_parsed:
        return jsonify({
            "status": "failed",
            "error_at": "agent1",
            "agent1": {"input": user_input, "output": req_parsed, "error": True},
            "source": source,
        })

    # Agent 2: Strategic Sourcing
    sourcing_json = source_all_items(req_json)
    sourcing_parsed = json.loads(sourcing_json)

    if "error" in sourcing_parsed:
        return jsonify({
            "status": "failed",
            "error_at": "agent2",
            "agent1": {"input": user_input, "output": req_parsed},
            "agent2": {"input": req_parsed, "output": sourcing_parsed, "error": True},
            "source": source,
        })

    # Pass intent_type through to Agent 3 for correction workflow
    intent_type = req_parsed.get("intent_type", "NEW_ORDER")

    # Agent 3: Email Drafts
    drafts_json = generate_all_drafts(sourcing_json, user_input, intent_type)
    drafts_parsed = json.loads(drafts_json)

    if "error" in drafts_parsed:
        return jsonify({
            "status": "failed",
            "error_at": "agent3",
            "agent1": {"input": user_input, "output": req_parsed},
            "agent2": {"input": req_parsed, "output": sourcing_parsed},
            "agent3": {"input": sourcing_parsed, "output": drafts_parsed, "error": True},
            "source": source,
        })

    return jsonify({
        "status": "success",
        "intent_type": intent_type,
        "agent1": {"input": user_input, "output": req_parsed},
        "agent2": {"input": req_parsed, "output": sourcing_parsed},
        "agent3": {"input": sourcing_parsed, "output": drafts_parsed},
        "drafts_json": drafts_json,
        "original_request": user_input,
        "source": source,
    })


@app.route("/run", methods=["POST"])
def run_pipeline():
    """Run the full 3-agent pipeline from text input."""
    data = request.get_json(force=True)
    user_input = data.get("request", "").strip()
    return _run_pipeline_from_text(user_input, "text")


@app.route("/extract", methods=["POST"])
def extract_and_run():
    """
    Accept an uploaded document (.docx, .pdf, .txt, image),
    extract text server-side, then run the pipeline.
    """
    if "file" not in request.files:
        return jsonify({"error": "no file uploaded"}), 400

    f = request.files["file"]
    if not f.filename:
        return jsonify({"error": "no filename provided"}), 400

    text, source_type = _extract_text_from_upload(f)

    if not text:
        return jsonify({
            "error": f"Could not extract text from {f.filename}. "
                     f"Source type: {source_type}. "
                     f"Try pasting the text directly.",
        }), 400

    return _run_pipeline_from_text(text, source_type)


@app.route("/send-emails", methods=["POST"])
def send_emails():
    """Send all drafted emails via SMTP."""
    data = request.get_json(force=True)
    drafts_json = data.get("drafts_json", "")

    if not drafts_json:
        return jsonify({"error": "missing drafts data"}), 400

    result_json = send_all_emails(drafts_json)
    result = json.loads(result_json)
    return jsonify(result)


# ── CSV Data API ──────────────────────────────────────────────────────

@app.route("/csv-data")
def csv_data():
    rows = _read_csv()
    return jsonify({"headers": list(rows[0].keys()) if rows else [], "rows": rows})


@app.route("/upload-csv", methods=["POST"])
def upload_csv():
    if "file" not in request.files:
        return jsonify({"error": "no file uploaded"}), 400

    f = request.files["file"]
    if not f.filename.endswith(".csv"):
        return jsonify({"error": "only .csv files accepted"}), 400

    f.save(CSV_PATH)
    rows = _read_csv()
    return jsonify({"headers": list(rows[0].keys()) if rows else [], "rows": rows})


def _read_csv() -> list[dict]:
    if not os.path.exists(CSV_PATH):
        return []
    with open(CSV_PATH, newline="", encoding="utf-8") as f:
        reader = csv.DictReader(f)
        return list(reader)


if __name__ == "__main__":
    app.run(host="0.0.0.0", port=5000, debug=True)